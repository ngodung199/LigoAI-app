import streamlit as st
import json
import uuid
import numpy as np
import string
from rank_bm25 import BM25Okapi
from groq import Groq
from docx import Document
from io import BytesIO
from supabase import create_client, Client

# ================== 1. KẾT NỐI ĐÁM MÂY ==================
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]

@st.cache_resource
def init_supabase():
 try:
    return create_client(SUPABASE_URL, SUPABASE_KEY)
 except:
    return None
     
supabase = init_supabase()


# ================== 2. BỘ MÁY TÌM KIẾM TỪ KHÓA (BM25) ==================
@st.cache_data
def load_laws():
 try:
    with open("legal_data.json", "r", encoding="utf-8") as f:
        return json.load(f)
 except:
    return []


def tokenize(text):
 if not text: return []
 text = text.lower()
 for p in string.punctuation:
    text = text.replace(p, ' ')
 return text.split()


@st.cache_resource
def get_bm25_index(_laws):
  if not _laws: return None
  corpus = [tokenize(item.get("content", "") + " " + item.get("title", "")) for item in _laws]
  return BM25Okapi(corpus)


laws = load_laws()
bm25_index = get_bm25_index(laws)


def retrieve_law_bm25(query, top_k=2):
 if bm25_index is None or not laws: return []
 scores = bm25_index.get_scores(tokenize(query))
 top_indices = np.argsort(scores)[::-1][:top_k]
 return [laws[i] for i in top_indices if scores[i] > 0]


# ================== 3. GIAO DIỆN CHUYÊN NGHIỆP ==================
st.set_page_config(page_title="LigoAI | Tư vấn Pháp lý", layout="wide")
st.markdown("""<style>.stButton button { border-radius: 8px; }</style>""", unsafe_allow_html=True)

if "conversations" not in st.session_state:
 uid = str(uuid.uuid4())
 st.session_state.conversations = {uid: []}
 st.session_state.current_chat = uid

current_chat_id = st.session_state.current_chat
current_messages = st.session_state.conversations[current_chat_id]

# --- SIDEBAR QUẢN LÝ ---
with st.sidebar:
 st.markdown("### LigoAI Legal")
 GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
 if st.button("➕ Cuộc hội thoại mới", use_container_width=True):
    new_id = str(uuid.uuid4())
    st.session_state.conversations[new_id] = []
    st.session_state.current_chat = new_id
    st.rerun()
  # Chèn đoạn này vào vị trí cuối cùng trong khối "with st.sidebar:"
    st.markdown("---")
    with st.expander("🛠️ Dành cho Ban giám khảo (Dữ liệu Admin)"):
        if st.button("Tải dữ liệu từ Supabase", use_container_width=True):
            try:
                # Kéo toàn bộ dữ liệu từ bảng chat_history về
                response = supabase.table("chat_history").select("*").execute()
                data = response.data
                
                if data:
                    # Hiển thị dưới dạng bảng cực kỳ chuyên nghiệp
                    st.dataframe(data, use_container_width=True)
                    st.caption(f"Tổng cộng: {len(data)} lượt truy vấn.")
                else:
                    st.info("Chưa có dữ liệu nào.")
            except Exception as e:
                st.error("Không thể kết nối máy chủ.")

 st.markdown("---")
 st.markdown("### Tiện ích văn bản")
 if st.button("Trích xuất Giấy đăng ký HKD", use_container_width=True):
    if len(current_messages) < 2:
        st.warning("Hãy trò chuyện để cung cấp thông tin trước.")
    else:
        with st.spinner("Đang soạn thảo file Word..."):
            try:
                client_tool = Groq(api_key=GROQ_API_KEY)
                chat_text = "\n".join([f"{m['role']}: {m['content']}" for m in current_messages])
                res = client_tool.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user",
                               "content": f"Trích xuất JSON (TEN_KHACH_HANG, DIA_CHI, NGANH_NGHE, VON_KINH_DOANH) từ: {chat_text}"}],
                    temperature=0.1
                )
                raw_text = res.choices[0].message.content
                data = json.loads(raw_text[raw_text.find('{'):raw_text.rfind('}') + 1])

                doc = Document()
                doc.add_heading('GIẤY ĐỀ NGHỊ ĐĂNG KÝ HỘ KINH DOANH', 0)
                doc.add_paragraph(f"Tên chủ hộ: {data.get('TEN_KHACH_HANG', '................')}")
                doc.add_paragraph(f"Địa chỉ: {data.get('DIA_CHI', '................')}")
                doc.add_paragraph(f"Ngành nghề: {data.get('NGANH_NGHE', '................')}")
                doc.add_paragraph(f"Vốn: {data.get('VON_KINH_DOANH', '................')}")

                bio = BytesIO()
                doc.save(bio)
                st.download_button("📥 Tải file Word", bio.getvalue(), "Dang_Ky_HKD.docx", type="primary")
            except:
                st.error("Chưa đủ thông tin để tạo đơn.")

# --- KHU VỰC CHAT CHÍNH ---
st.markdown("<h3 style='text-align: center;'>Xin chào, tôi là LigoAI</h3>", unsafe_allow_html=True)

suggestion_clicked = None
c1, c2 = st.columns(2)
with c1:
  if st.button("Mở tiệm tạp hóa doanh thu 150tr thì đóng thuế gì?",
             use_container_width=True): suggestion_clicked = "Mở tiệm tạp hóa doanh thu 150tr thì đóng thuế gì?"
with c2:
  if st.button("Thủ tục đăng ký hộ kinh doanh cần giấy tờ gì?",
             use_container_width=True): suggestion_clicked = "Thủ tục đăng ký hộ kinh doanh cần giấy tờ gì?"

for msg in current_messages:
  with st.chat_message(msg["role"]):
    st.markdown(msg["content"])
    if msg["role"] == "assistant" and msg.get("retrieved"):
        with st.expander("📑 Căn cứ pháp lý"):
            for item in msg["retrieved"]:
                st.markdown(f"**{item.get('title', '')}**\n*{item.get('content', '')}*")

user_input = st.chat_input("Nhập vấn đề pháp lý của bạn tại đây...")
prompt = user_input or suggestion_clicked

# BỨC TƯỜNG LỬA: Chỉ chạy khi prompt có nội dung thật sự, chặn đứng chuỗi rỗng và chữ "None"
if prompt and str(prompt).strip() != "" and str(prompt).strip() != "None":

# 1. Đẩy dữ liệu lên Supabase
 if supabase:
    try:
        supabase.table("chat_history").insert({"session_id": current_chat_id, "user_query": prompt}).execute()
    except:
        pass  # Lỗi mạng bỏ qua, web vẫn chạy tiếp

# 2. Lưu và hiển thị câu hỏi
 st.session_state.conversations[current_chat_id].append({"role": "user", "content": prompt})
 with st.chat_message("user"):
    st.markdown(prompt)

# 3. AI suy nghĩ và phản hồi
 with st.chat_message("assistant"):
    msg_placeholder = st.empty()
    full_res = ""
    retrieved = retrieve_law_bm25(prompt)

    if not retrieved:
        full_res = "LigoAI chưa tìm thấy quy định phù hợp trong hệ thống."
        msg_placeholder.markdown(full_res)
    else:
        client = Groq(api_key=GROQ_API_KEY)
        context = "\n".join([f"- {i['title']}: {i['content']}" for i in retrieved])
        sys_prompt = "Bạn là trợ lý luật pháp chuyên nghiệp. Trả lời dựa trên ngữ cảnh được cung cấp."
        user_msg = f"Ngữ cảnh:\n{context}\n\nCâu hỏi: {prompt}"

        try:
            stream = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": user_msg}],
                stream=True
            )
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    full_res += chunk.choices[0].delta.content
                    msg_placeholder.markdown(full_res + "▌")
            msg_placeholder.markdown(full_res)
        except:
            full_res = "⚠️ Máy chủ AI đang bận."
            msg_placeholder.markdown(full_res)

    st.session_state.conversations[current_chat_id].append(
        {"role": "assistant", "content": full_res, "retrieved": retrieved})

# ĐÃ XÓA LỆNH st.rerun() GÂY LỖI Ở ĐÂY



